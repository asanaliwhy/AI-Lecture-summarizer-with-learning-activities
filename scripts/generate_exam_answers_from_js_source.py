from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE_PATH = Path(r"C:\Users\NURMA\OneDrive\Рабочий стол\projects\DIPLOMA\scripts\exam_tickets_source.js")
OUTPUT_PATH = Path(r"C:\Users\NURMA\Downloads\Орталықтандырылған_жылумен_жабдықтау_жүйелері_жауаптар.docx")


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def parse_tickets(raw: str) -> list[dict[str, object]]:
    ticket_positions = list(re.finditer(r"num:\s*(\d+)", raw))
    if len(ticket_positions) != 25:
        raise RuntimeError(f"Expected 25 tickets, found {len(ticket_positions)}")

    qa_pattern = re.compile(r'q:\s*"(.*?)"\s*,\s*a:\s*`(.*?)`', re.S)
    tickets: list[dict[str, object]] = []

    for index, match in enumerate(ticket_positions):
        ticket_num = int(match.group(1))
        start = match.start()
        end = ticket_positions[index + 1].start() if index + 1 < len(ticket_positions) else len(raw)
        section = raw[start:end]

        qa_matches = qa_pattern.findall(section)
        if len(qa_matches) != 3:
            raise RuntimeError(f"Ticket {ticket_num} expected 3 questions, found {len(qa_matches)}")

        tickets.append(
            {
                "num": ticket_num,
                "qs": [
                    {"q": normalize_text(question), "a": normalize_text(answer)}
                    for question, answer in qa_matches
                ],
            }
        )

    return tickets


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def build_document(tickets: list[dict[str, object]]) -> Document:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Орталықтандырылған жылумен жабдықтау жүйелері\nЕмтихан билеттеріне толық жауаптар")
    title_run.bold = True
    title_run.font.size = Pt(15)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Пайдаланушы ұсынған жауаптар негізінде жаңартылған нұсқа")

    for ticket in tickets:
        document.add_paragraph()
        heading = document.add_paragraph()
        heading_run = heading.add_run(f"ЕМТИХАН БИЛЕТ № {ticket['num']}")
        heading_run.bold = True
        heading_run.font.size = Pt(14)

        for qa in ticket["qs"]:  # type: ignore[index]
            question = qa["q"]  # type: ignore[index]
            answer = qa["a"]  # type: ignore[index]

            question_paragraph = document.add_paragraph()
            question_run = question_paragraph.add_run(str(question))
            question_run.bold = True

            answer_paragraph = document.add_paragraph(str(answer))
            answer_paragraph.paragraph_format.space_after = Pt(10)

        if ticket["num"] != 25:
            document.add_page_break()

    return document


def main() -> None:
    raw = SOURCE_PATH.read_text(encoding="utf-8")
    tickets = parse_tickets(raw)
    tickets.sort(key=lambda item: int(item["num"]))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(tickets)
    document.save(OUTPUT_PATH)
    print("DOCX_CREATED_FROM_USER_SOURCE")


if __name__ == "__main__":
    main()
