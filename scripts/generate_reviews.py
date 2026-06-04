"""
Generate two REVIEW documents in DOCX format — one page each, no technical jargon, 3 paragraphs each.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TITLE = "Development of an AI-based Platform for Lecture Summaries and Learning Activities"
AUTHORS = "Nurmagambet Asanali, Bayadilov Asanali"
SUPERVISOR = "Kozhakhmet Z."
PROGRAM = '6B06102-Software Engineering'

OUTPUT_DIR = r"C:\Users\NURMA\OneDrive\Рабочий стол"


def make_run(paragraph, text, font_name="Times New Roman", size=12):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    return run


def add_centered(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    make_run(p, text, size=size)
    return p


def add_justified(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(1.25)
    make_run(p, text, size=size)
    return p


def add_left(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    make_run(p, text, size=size)
    return p


def add_empty(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(6)
    return p


def build_review(paragraphs, reviewer_name, reviewer_title, reviewer_dept, reviewer_univ, reviewer_city):
    doc = Document()

    # Slightly tighter margins to ensure it all fits on one page with 3 paragraphs
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    add_centered(doc, "REVIEW")
    add_centered(doc, "for diploma project of")
    add_centered(doc, AUTHORS)
    add_centered(doc, "on the topic")
    add_centered(doc, f'"{TITLE}"')
    add_centered(doc, f"Research supervisor: {SUPERVISOR}")
    add_empty(doc)

    for paragraph in paragraphs:
        add_justified(doc, paragraph)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(1.25)
    make_run(p, f'The reviewed diploma project meets the qualification requirements, and the authors are awarded the academic degree Bachelor in Information and Communication Technologies in the educational program "{PROGRAM}" with an \u00ab___\u00bb grade.')

    add_empty(doc)

    add_left(doc, f"Reviewer: {reviewer_name}")
    add_left(doc, reviewer_title)
    add_left(doc, reviewer_dept)
    add_left(doc, reviewer_univ)
    add_left(doc, reviewer_city)
    add_empty(doc)
    add_left(doc, "\u00ab____\u00bb  ___________ 2026")

    return doc


# ── Review 1 ──
review1_paragraphs = [
    "The diploma project focuses on the development of a web platform called Lectura, designed to help students process lecture materials more efficiently. The system addresses a real and practical problem in modern education \u2014 the difficulty of converting large volumes of video and text content into useful study materials. By allowing users to paste a video link or upload a document, the platform automates the creation of study aids, saving significant time and effort for learners who would otherwise have to manually synthesize this information.",
    
    "The proposed solution is well thought out and covers the full cycle from content input to active learning. Once a file or link is provided, the system automatically generates structured summaries in several formats, quizzes with different difficulty levels, flashcard decks for memorization practice, and ready-to-use presentation slides. The platform provides a personal dashboard where students can track their progress, view activity statistics, and manage all generated resources in a unified library. Additionally, the inclusion of an AI-powered chat feature that answers questions based strictly on the original lecture content adds significant value by providing a reliable virtual assistant for studying.",
    
    "The project is clearly structured, logically organized, and demonstrates a consistent approach to solving the stated problem. The authors have shown the ability to identify a meaningful educational challenge, design an appropriate solution, and implement it as a fully working product. The final platform is user-friendly, highly functional, and ready for practical use. Overall, the diploma project is relevant, practically significant, and well-executed, proving that the authors fully meet the requirements for their academic degree."
]

review1 = build_review(
    review1_paragraphs,
    reviewer_name="Suleimenova Zarina",
    reviewer_title="PhD, Associate Professor",
    reviewer_dept="Department of Computer Science",
    reviewer_univ="Astana IT University",
    reviewer_city="Astana, Kazakhstan",
)
review1_path = os.path.join(OUTPUT_DIR, "REVIEW_1.docx")
review1.save(review1_path)
print(f"Saved: {review1_path}")


# ── Review 2 ──
review2_paragraphs = [
    "The diploma project presents the development of Lectura, an intelligent educational platform that transforms lecture recordings and documents into a set of interactive learning resources. The topic is highly relevant given the growing need for digital tools that support independent and efficient learning in academic environments. The authors recognized that students often struggle to organize their notes and test their knowledge effectively, and set out to build a comprehensive system that takes raw educational content and turns it into actionable study materials without requiring complex manual input.",
    
    "The platform is designed with a clear focus on usability and practical application. After a user submits a video link or uploads a file, the system automatically produces detailed summaries, assessment quizzes, flashcards with a built-in review schedule, and presentation slides that can be edited and exported. All generated materials are accessible through a single library with search and filtering capabilities. Furthermore, a personal dashboard provides an overview of study activity, including streaks and weekly goals. The ability to ask questions about the lecture content through an AI assistant that stays grounded in the source material is a noteworthy addition that helps prevent the spread of incorrect information.",
    
    "The project demonstrates a well-organized workflow from initial problem definition to a functioning and polished end product. The authors have shown a strong understanding of the educational challenges they set out to address and have delivered a cohesive solution that covers content processing, knowledge assessment, and long-term retention support. The resulting platform is intuitive and offers a genuine benefit to its target audience. The project is relevant, well-motivated, and fully meets the expected standards for a diploma work at the bachelor level."
]

review2 = build_review(
    review2_paragraphs,
    reviewer_name="Tlegenov Daulet",
    reviewer_title="MSc, Senior Lecturer",
    reviewer_dept="Department of Software Engineering",
    reviewer_univ="L.N. Gumilyov Eurasian National University",
    reviewer_city="Astana, Kazakhstan",
)
review2_path = os.path.join(OUTPUT_DIR, "REVIEW_2.docx")
review2.save(review2_path)
print(f"Saved: {review2_path}")
