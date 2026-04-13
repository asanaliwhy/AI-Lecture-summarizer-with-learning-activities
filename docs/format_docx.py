"""
Professional DOCX Formatter for AITU Diploma Thesis
Applies academic formatting standards to diploma-work.docx
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

INPUT_FILE = r"C:\Users\NURMA\Downloads\diploma-work.docx"
OUTPUT_FILE = r"C:\Users\NURMA\Downloads\diploma-work-formatted.docx"

def set_cell_shading(cell, color):
    """Set the background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color="000000", size="4"):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def format_run(run, font_name="Times New Roman", font_size=14, bold=False, italic=False, color=None):
    """Apply formatting to a text run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}" w:cs="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)

def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0,
                         space_after=0, line_spacing=1.5, first_line_indent=None,
                         keep_together=False, keep_with_next=False, page_break_before=False):
    """Apply formatting to a paragraph."""
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    else:
        pf.first_line_indent = None
    pf.keep_together = keep_together
    pf.keep_with_next = keep_with_next
    pf.page_break_before = page_break_before

def is_section_heading(text):
    """Check if text is a main section heading (e.g., '1. INTRODUCTION')."""
    text = text.strip()
    # Match patterns like "1. INTRODUCTION" or "INTRODUCTION" or "1. Introduction"
    patterns = [
        r'^\d+\.\s+[A-Z]',                # "1. INTRODUCTION"
        r'^ABSTRACT',                       # "ABSTRACT"
        r'^TABLE OF CONTENTS',              # "TABLE OF CONTENTS"
        r'^REFERENCES',                     # "REFERENCES"
        r'^APPENDIX',                       # "APPENDIX"
    ]
    for p in patterns:
        if re.match(p, text):
            return True
    return False

def is_subsection_heading(text):
    """Check if text is a subsection heading (e.g., '1.1. Background')."""
    text = text.strip()
    return bool(re.match(r'^\d+\.\d+\.?\s+', text))

def is_table_caption(text):
    """Check if text is a table caption."""
    text = text.strip()
    return text.startswith("Table ") and "—" in text

def is_figure_caption(text):
    """Check if text is a figure caption."""
    text = text.strip()
    return text.startswith("Figure ") and "—" in text

def is_list_item(text):
    """Check if text starts with a list marker."""
    text = text.strip()
    return bool(re.match(r'^[\-•]\s', text)) or bool(re.match(r'^\d+[\.\)]\s', text))

def main():
    print(f"Opening: {INPUT_FILE}")
    doc = Document(INPUT_FILE)
    
    # ──────────────────────────────────────────────────────────────
    # 1. PAGE SETUP
    # ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)  # A4
        section.left_margin = Cm(3.0)   # 30mm binding margin
        section.right_margin = Cm(1.5)  # 15mm
        section.top_margin = Cm(2.0)    # 20mm
        section.bottom_margin = Cm(2.0) # 20mm
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
    
    print("✓ Page setup: A4, margins 30/15/20/20mm")
    
    # ──────────────────────────────────────────────────────────────
    # 2. ADD PAGE NUMBERS IN FOOTER
    # ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing footer
        for p in footer.paragraphs:
            p.clear()
        
        # Add page number
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add PAGE field
        run = fp.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        
        run3 = fp.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)
        
        for r in [run, run2, run3]:
            format_run(r, font_size=12)
    
    print("✓ Page numbers added to footer")
    
    # ──────────────────────────────────────────────────────────────
    # 3. FORMAT PARAGRAPHS
    # ──────────────────────────────────────────────────────────────
    cover_page_end = 29  # Index of last cover page paragraph
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Skip empty paragraphs
        if not text:
            # Make empty paragraphs small
            for run in para.runs:
                format_run(run, font_size=14)
            set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.5)
            continue
        
        # ── COVER PAGE (paragraphs 0-29) ──
        if i <= cover_page_end:
            for run in para.runs:
                format_run(run, font_size=14, bold=True)
            
            if i == 0:  # "Astana IT University"
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=36, space_after=6, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=16, bold=True)
            elif i in [7, 8]:  # Authors
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=3, space_after=3, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=14, bold=False)
            elif "Development AI" in text or "Summaries and Learning" in text:  # Title
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=12, space_after=6, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=18, bold=True)
            elif "6B06101" in text:  # Specialty
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=6, space_after=6, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=14, bold=True)
            elif "Diploma project" in text:
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=12, space_after=12, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=16, bold=True)
            elif "Supervisor" in text or "Kozhakhmet" in text or "School of" in text:
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    space_before=3, space_after=3, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=14, bold=False)
            elif "Kazakhstan" in text or "Astana, 2026" in text:
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=24, space_after=6, line_spacing=1.5)
                for run in para.runs:
                    format_run(run, font_size=14, bold=False)
            else:
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=3, space_after=3, line_spacing=1.5)
            continue
        
        # ── MAIN SECTION HEADINGS ──
        if is_section_heading(text):
            set_paragraph_format(para,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_before=24, space_after=18,
                                line_spacing=1.5,
                                first_line_indent=0,
                                page_break_before=True,
                                keep_with_next=True)
            for run in para.runs:
                format_run(run, font_size=18, bold=True)
                run.font.all_caps = True
            continue
        
        # ── SUBSECTION HEADINGS ──
        if is_subsection_heading(text):
            set_paragraph_format(para,
                                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                space_before=18, space_after=12,
                                line_spacing=1.5,
                                first_line_indent=0,
                                keep_with_next=True)
            for run in para.runs:
                format_run(run, font_size=16, bold=True)
            continue
        
        # ── TABLE CAPTION ──
        if is_table_caption(text):
            set_paragraph_format(para,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_before=12, space_after=6,
                                line_spacing=1.5,
                                first_line_indent=0,
                                keep_with_next=True)
            for run in para.runs:
                format_run(run, font_size=12, bold=True, italic=True)
            continue
        
        # ── FIGURE CAPTION ──
        if is_figure_caption(text):
            set_paragraph_format(para,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_before=6, space_after=12,
                                line_spacing=1.5,
                                first_line_indent=0)
            for run in para.runs:
                format_run(run, font_size=12, bold=True, italic=True)
            continue
        
        # ── LIST ITEMS ──
        if para.style.name == "List Paragraph" or is_list_item(text):
            set_paragraph_format(para,
                                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                space_before=3, space_after=3,
                                line_spacing=1.5,
                                first_line_indent=0)
            # Set left indent for list items
            para.paragraph_format.left_indent = Cm(1.25)
            for run in para.runs:
                format_run(run, font_size=14)
            continue
        
        # ── REFERENCES ──
        if text.startswith("[") and re.match(r'^\[\d+\]', text):
            set_paragraph_format(para,
                                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                space_before=3, space_after=3,
                                line_spacing=1.5,
                                first_line_indent=0)
            para.paragraph_format.left_indent = Cm(1.25)
            para.paragraph_format.first_line_indent = Cm(-1.25)  # Hanging indent
            for run in para.runs:
                format_run(run, font_size=12)
            continue
        
        # ── REGULAR BODY TEXT ──
        set_paragraph_format(para,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=0, space_after=6,
                            line_spacing=1.5,
                            first_line_indent=1.25)
        for run in para.runs:
            format_run(run, font_size=14)
    
    print("✓ All paragraphs formatted")
    
    # ──────────────────────────────────────────────────────────────
    # 4. FORMAT TABLES
    # ──────────────────────────────────────────────────────────────
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Style cell borders
                set_cell_borders(cell, color="333333", size="4")
                
                # Header row styling
                if row_idx == 0:
                    set_cell_shading(cell, "E8EDF2")
                
                # Format cell text
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pf = para.paragraph_format
                    pf.space_before = Pt(3)
                    pf.space_after = Pt(3)
                    pf.line_spacing = 1.15
                    
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                        if row_idx == 0:
                            run.font.bold = True
                            run.font.size = Pt(11)
    
    print("✓ Tables formatted")
    
    # ──────────────────────────────────────────────────────────────
    # 5. SAVE
    # ──────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"\n✅ Saved to: {OUTPUT_FILE}")
    print("Done! Open the file in Word to verify.")

if __name__ == "__main__":
    main()
